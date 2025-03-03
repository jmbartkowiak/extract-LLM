# resume_builder.py
# v1.0.3
# 2-27-25


"""
Resume Builder Module

This module assembles the final resume by injecting optimized content into a DOCX template.
Flow:
1. Load the DOCX template from STATIC_DATA.
2. Validate and parse optimized content into an OptimizedContent data structure.
3. Extract placeholders from the template.
4. Replace placeholders with the optimized objective, skills, and bullet points.
5. Save the final resume and log its details.
"""

import os
import json
import datetime
from pathlib import Path
from typing import Dict, List, Optional, Any, Union
from dataclasses import dataclass

from docx import Document
from logging_manager import log_process
from placeholder_matcher import PlaceholderMatcher
from config_manager import CONFIG
from helpers import safe_file_write

@dataclass
class OptimizedContent:
    """
    Container for optimized resume content.
    Expected fields:
      - objective: The refined overview statement.
      - skills: A list of exactly 10 skills.
      - bullets: A list of bullet point dictionaries.
      - match_rating: The match rating percentage.
      - job_data: A dictionary with job metadata.
    """
    objective: str
    skills: List[str]
    bullets: List[Dict[str, str]]
    match_rating: float
    job_data: Dict[str, str]
    
    @classmethod
    def from_dict(cls, data: Dict[str, Any]) -> 'OptimizedContent':
        return cls(
            objective=data.get('new_objective', ''),
            skills=data.get('optimized_skills', '').split(', '),
            bullets=data.get('optimized_bullets', {}).get('bullets', []),
            match_rating=data.get('job_match_evaluation', {}).get('match_rating', 0),
            job_data=data.get('job_json', {})
        )
    
    def validate(self) -> None:
        """
        Validates that the optimized content meets the requirements.
        Ensures a non-empty objective and exactly 10 skills.
        """
        if not self.objective:
            raise ValueError("Missing objective statement")
        if len(self.skills) != 10:
            raise ValueError(f"Expected 10 skills, got {len(self.skills)}")
        required_fields = ['jid', 'Company Name', 'Title']
        missing = [field for field in required_fields if not self.job_data.get(field)]
        if missing:
            raise ValueError(f"Missing required job fields: {', '.join(missing)}")

class ResumeBuilderError(Exception):
    """Custom exception for resume building errors."""
    pass

def get_template_path() -> Path:
    """
    Returns the path to the resume DOCX template.
    """
    template_path = Path(CONFIG.get("STATIC_DATA_DIR", "STATIC_DATA")) / "resume_template" / "template-resume.docx"
    if not template_path.exists():
        raise ResumeBuilderError(f"Template not found at {template_path}")
    return template_path

def create_output_directory(job_data: Dict[str, str], base_dir: Optional[str] = None) -> Path:
    """
    Creates an output directory based on the job's metadata.
    """
    if base_dir is None:
        base_dir = CONFIG.get("FINISHED_JOB_RESUME_DIR", "FINISHED_JOB_RESUME")
    try:
        today = datetime.datetime.now().strftime("%Y%m%d")
        company = job_data.get("Company Name", "Unknown").replace(" ", "")
        title = job_data.get("Title", "Unknown").replace(" ", "")
        jid = job_data.get("jid", "Unknown")
        output_dir = Path(base_dir) / f"{today}_{company}_{title}_{jid}"
        output_dir.mkdir(parents=True, exist_ok=True)
        return output_dir
    except Exception as e:
        raise ResumeBuilderError(f"Failed to create output directory: {e}")

def replace_placeholder_text(paragraph: Any, placeholder: str, new_text: str, preserve_formatting: bool = True) -> None:
    """
    Replaces the placeholder text within a paragraph.
    If preserve_formatting is True, it replaces text in each run.
    """
    if not preserve_formatting:
        paragraph.text = paragraph.text.replace(placeholder, new_text)
        return
    for run in paragraph.runs:
        if placeholder in run.text:
            run.text = run.text.replace(placeholder, new_text)

def inject_content(doc: Document, content: OptimizedContent, skill_placeholders: List[str]) -> None:
    """
    Injects optimized content into the document by replacing placeholders.
    
    Flow:
    - Replace <OverView> with the objective.
    - For each skill placeholder (<SKILL N>), replace with the corresponding skill.
    - Replace bullet placeholders with bullet text.
    """
    try:
        # Replace overview placeholder.
        for para in doc.paragraphs:
            if "<OverView>" in para.text:
                replace_placeholder_text(para, "<OverView>", content.objective)
        # Replace skill placeholders.
        for i, skill in enumerate(content.skills, 1):
            placeholder = f"<SKILL {i}>"
            for para in doc.paragraphs:
                if placeholder in para.text:
                    replace_placeholder_text(para, placeholder, skill)
        # Replace bullet placeholders.
        for i, bullet in enumerate(content.bullets, 1):
            overview_placeholder = f"<Experience-Bullet{i}-BoldedOverview-J1>"
            detail_placeholder = f"<Experience-Bullet{i}-J1>"
            for para in doc.paragraphs:
                if overview_placeholder in para.text:
                    replace_placeholder_text(para, overview_placeholder, bullet.get("bolded_overview", ""))
                if detail_placeholder in para.text:
                    replace_placeholder_text(para, detail_placeholder, bullet.get("description", ""))
    except Exception as e:
        raise ResumeBuilderError(f"Failed to inject content: {e}")

def log_resume_details(content: OptimizedContent, output_path: Path) -> None:
    """
    Logs details of the generated resume for analytics.
    """
    job_data = content.job_data
    details = {
        "job_description_jid": job_data.get("jid", "Unknown"),
        "final_resume_date": datetime.datetime.now().strftime("%Y-%m-%d"),
        "final_resume_match_percentage": content.match_rating,
        "final_resume_filename": output_path.name,
        "final_resume_link": str(output_path.resolve())
    }
    log_process(f"Resume details: {json.dumps(details)}", "INFO", module="ResumeBuilder")

def build_final_resume(optimized_data: Dict[str, Any], output_dir: Optional[str] = None) -> Optional[Path]:
    """
    Builds the final resume document.
    
    Steps:
    1. Convert optimized_data to an OptimizedContent instance and validate it.
    2. Load the DOCX template.
    3. Extract skill placeholders from the document.
    4. Inject optimized content into the template.
    5. Save the final document and log its details.
    """
    try:
        content = OptimizedContent.from_dict(optimized_data)
        content.validate()
        template_path = get_template_path()
        out_dir = create_output_directory(content.job_data, output_dir)
        output_path = out_dir / "final_resume.docx"
        doc = Document(template_path)
        # Extract skill placeholders from the document text.
        all_text = "\n".join([p.text for p in doc.paragraphs])
        skill_placeholders = list(PlaceholderMatcher.extract_skill_placeholders(all_text).values())
        inject_content(doc, content, skill_placeholders)
        doc.save(output_path)
        log_process(f"Final resume saved at {output_path}", "INFO", module="ResumeBuilder")
        log_resume_details(content, output_path)
        return output_path
    except Exception as e:
        log_process(f"Failed to build resume: {e}", "ERROR", module="ResumeBuilder")
        raise ResumeBuilderError(e)

# End of resume_builder.py
