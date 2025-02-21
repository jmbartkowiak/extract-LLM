"""
resume_extractor.py
2/20/2025
v8.0.4

This module processes resume files from various formats (PDF, DOCX, TXT),
extracts structured data via LLM API calls, and organizes the files into
appropriate directories.

Enhancements:
- Uses the 'resume_extraction_strict_prompt' to enforce valid JSON output.
- Sanitizes resume text prior to LLM calls.
- Enhanced JSON cleanup and fallback parsing.
- Extracts bullet points from DOCX tables and integrates them into the text.
- Adds a file processing check to skip resumes that have already been processed.
"""

import os
import json
import datetime
import shutil
import re
from pathlib import Path
from typing import Dict, List, Optional, Any, Union
from dataclasses import dataclass

import fitz  # PyMuPDF for PDF extraction
from docx import Document

from logging_manager import log_process, log_application
from config_manager import CONFIG
from helpers import validate_file_path, clean_filename

# Type aliases
JSON = Dict[str, Any]
FilePath = Union[str, Path]

@dataclass
class ResumeData:
    """Container for extracted resume data"""
    rid: str
    objective: str
    skills_list: List[str]
    jobs_section: List[Dict[str, Any]]
    education: List[Dict[str, str]]
    certifications: List[str]
    raw_text: str
    source_file: str
    extraction_date: str
    usage_count: int = 0

    @classmethod
    def from_dict(cls, data: Dict[str, Any]) -> 'ResumeData':
        return cls(
            rid=data.get('rid', ''),
            objective=data.get('objective', ''),
            skills_list=data.get('skills_list', []),
            jobs_section=data.get('jobs_section', []),
            education=data.get('education', []),
            certifications=data.get('certifications', []),
            raw_text=data.get('raw_text', ''),
            source_file=data.get('source_file', ''),
            extraction_date=data.get('extraction_date', ''),
            usage_count=data.get('usage_count', 0)
        )

class ResumeExtractionError(Exception):
    """Custom exception for resume extraction errors"""
    pass

def is_resume_file_processed(resume_file: Path) -> bool:
    """
    Checks if a resume file has already been processed by scanning the EXTRACTED_DATA/resume_analysis folder.
    It compares the "source_file" field in each JSON file to the resume file's name.
    """
    extracted_dir = Path(CONFIG["EXTRACTED_DATA_DIR"]) / "resume_analysis"
    if not extracted_dir.exists():
        return False
    for json_file in extracted_dir.glob("RES-*.json"):
        try:
            with open(json_file, "r", encoding="utf-8") as f:
                data = json.load(f)
                if data.get("source_file", "") == resume_file.name:
                    return True
        except Exception:
            continue
    return False

def sanitize_for_json(text: str) -> str:
    """
    Escapes backslashes and quotes that might break JSON structure.
    """
    text = text.replace("\\", "\\\\")
    text = text.replace("\"", "\\\"")
    return text

def load_prompts() -> Dict[str, Any]:
    """
    Loads all prompts from all_prompts.json.
    """
    prompts_path = Path(CONFIG["STATIC_DATA_DIR"]) / "prompt_templates" / "all_prompts.json"
    if not prompts_path.exists():
        raise FileNotFoundError(f"Prompt file not found at {prompts_path}")
    with open(prompts_path, "r", encoding="utf-8") as f:
        return json.load(f)

def extract_bullets_from_tables(doc: Document) -> List[Dict[str, str]]:
    """
    Extracts bullet points from tables in a DOCX document.
    Assumes each row in a table has at least two cells:
      - The first cell contains the bolded bullet title.
      - The second cell contains the bullet description.
    Returns a list of dictionaries with keys "bolded_overview" and "description".
    """
    bullets = []
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) >= 2:
                bullet_title = row.cells[0].text.strip()
                bullet_desc = row.cells[1].text.strip()
                if bullet_title or bullet_desc:
                    bullets.append({
                        "bolded_overview": bullet_title,
                        "description": bullet_desc
                    })
    return bullets

def extract_text_from_file(file_path: FilePath) -> Optional[str]:
    file_path = Path(file_path)
    try:
        if file_path.suffix.lower() == ".txt":
            return file_path.read_text(encoding="utf-8")
        elif file_path.suffix.lower() == ".pdf":
            text = []
            with fitz.open(str(file_path)) as doc:
                for page in doc:
                    text.append(page.get_text("text"))
            return "\n".join(text)
        elif file_path.suffix.lower() == ".docx":
            doc = Document(file_path)
            paragraph_text = "\n".join(para.text for para in doc.paragraphs if para.text.strip() != "")
            table_bullets = extract_bullets_from_tables(doc)
            bullet_text = "\n".join(
                f"• {b['bolded_overview']}: {b['description']}" for b in table_bullets
            )
            combined_text = "\n".join([paragraph_text, bullet_text]).strip()
            return combined_text
        else:
            raise ResumeExtractionError(f"Unsupported file type: {file_path.suffix}")
    except Exception as e:
        log_process(f"Error reading {file_path.name}: {e}", "ERROR")
        return None

def clean_api_response(response: Any) -> Dict[str, Any]:
    """
    Cleans API response and returns a dictionary with required fields.
    Enhanced with better JSON parsing and error handling.
    """
    # Log the initial raw response with type information
    log_process(f"Initial raw API response type: {type(response)}", "DEBUG", immediate=True)
    log_process(f"Initial raw API response: {repr(response)}", "DEBUG", immediate=True)
    
    default_data = {
        "objective": "",
        "skills_list": [],
        "jobs_section": [],
        "education": [],
        "certifications": []
    }
    
    try:
        # Handle LLMResponse object
        if hasattr(response, 'content'):
            response = response.content
            log_process(f"Extracted content from LLMResponse: {response}", "DEBUG", immediate=True)
        
        # First attempt: direct JSON parsing
        if isinstance(response, str):
            # Remove markdown blocks if present
            if "```json" in response:
                response = response.split("```json", 1)[1].split("```", 1)[0]
            elif "```" in response:
                response = response.split("```", 1)[1].split("```", 1)[0]
            
            # Find and extract JSON object
            start_idx = response.find('{')
            end_idx = response.rfind('}')
            if start_idx != -1 and end_idx != -1:
                response = response[start_idx:end_idx + 1]
            
            # Clean the response
            response = response.strip()
            # Remove any leading/trailing whitespace or newlines from each line
            lines = [line.strip() for line in response.splitlines()]
            # Join lines and ensure proper JSON formatting
            response = ''.join(lines)
            response = response.replace("'", '"')
            # Ensure we have proper JSON structure
            if not response.startswith('{'):
                response = '{' + response
            if not response.endswith('}'):
                response = response + '}'
            
            try:
                # First attempt: direct JSON parsing
                parsed = json.loads(response)
            except json.JSONDecodeError as e:
                log_process(f"First JSON parsing attempt failed: {e}", "DEBUG")
                try:
                    # Second attempt: try to fix common JSON formatting issues
                    response = re.sub(r',\s*}', '}', response)  # Remove trailing commas
                    response = re.sub(r',\s*]', ']', response)  # Remove trailing commas in arrays
                    response = re.sub(r'\s+', ' ', response)    # Normalize whitespace
                    parsed = json.loads(response)
                except json.JSONDecodeError as e2:
                    log_process(f"Second JSON parsing attempt failed: {e2}", "DEBUG")
                    # Return default data
                    parsed = default_data
        else:
            parsed = response
        
        # Ensure all required fields are present with correct types
        result = default_data.copy()
        if isinstance(parsed, dict):
            # Handle objective
            result["objective"] = str(parsed.get("objective", ""))
            
            # Handle skills_list
            skills = parsed.get("skills_list", [])
            if isinstance(skills, list):
                result["skills_list"] = [str(s) for s in skills]
            
            # Handle jobs_section
            jobs = parsed.get("jobs_section", [])
            if isinstance(jobs, list):
                result["jobs_section"] = []
                for job in jobs:
                    if isinstance(job, dict):
                        cleaned_job = {
                            "title": str(job.get("title", "")),
                            "company": str(job.get("company", "")),
                            "dates": str(job.get("dates", "")),
                            "bullets": []
                        }
                        bullets = job.get("bullets", [])
                        if isinstance(bullets, list):
                            for bullet in bullets:
                                if isinstance(bullet, dict):
                                    cleaned_bullet = {
                                        "bolded_overview": str(bullet.get("bolded_overview", "")),
                                        "description": str(bullet.get("description", ""))
                                    }
                                    cleaned_job["bullets"].append(cleaned_bullet)
                        result["jobs_section"].append(cleaned_job)
            
            # Handle education and certifications
            result["education"] = [str(e) for e in parsed.get("education", [])]
            result["certifications"] = [str(c) for c in parsed.get("certifications", [])]
        
        # Log the cleaned result
        log_process(f"Cleaned data: {json.dumps(result, indent=2)}", "DEBUG", immediate=True)
        
        return result
        
    except Exception as e:
        log_process(f"Failed to clean API response: {e}", "ERROR")
        return default_data

def extract_resume_data(raw_text: str, file_name: str) -> ResumeData:
    """
    Extracts structured resume data using the 'resume_extraction_strict_prompt'.
    Enhanced with fallback JSON parsing and additional logging.
    
    Args:
        raw_text: The raw text extracted from the resume file
        file_name: The name of the source file
        
    Returns:
        ResumeData object containing the structured data
        
    Raises:
        ResumeExtractionError: If extraction fails at any stage
    """
    from api_interface import call_api  # Avoid circular dependency
    try:
        sanitized_text = sanitize_for_json(raw_text)
        prompts = load_prompts()
        strict_prompt_template = prompts["resume_extraction_strict_prompt"]["prompt"]
        prompt = strict_prompt_template.format(resume_text=sanitized_text)
        
        # Call API with strict JSON requirement and enhanced system message
        result = call_api(
            prompt=prompt,
            system_message="""
            You MUST return a valid JSON object with EXACTLY these fields:
            {
              "objective": "string",
              "skills_list": ["string", "string", ...],
              "jobs_section": [
                {
                  "title": "string",
                  "company": "string",
                  "dates": "string",
                  "bullets": [
                    {
                      "bolded_overview": "string",
                      "description": "string"
                    }
                  ]
                }
              ],
              "education": ["string"],
              "certifications": ["string"]
            }
            
            IMPORTANT:
            1. Start response with '{' and end with '}'
            2. Use double quotes for all strings
            3. No comments or extra text
            4. No trailing commas
            5. No newlines in values
            """
        )
        if not result:
            raise ResumeExtractionError("Empty API response from LLM.")
        
        # Log the raw response for debugging
        log_process(f"Raw API response before cleaning: {repr(result)}", "DEBUG", immediate=True)
        
        # Try to parse the response directly first
        try:
            if isinstance(result, str):
                # Remove any leading/trailing whitespace and newlines
                result = result.strip()
                # Try to parse as JSON
                resume_data = json.loads(result)
            else:
                # If it's not a string, try to get the content
                if hasattr(result, 'content'):
                    result = result.content
                resume_data = clean_api_response(result)
            
            log_process(f"Parsed resume data: {json.dumps(resume_data, indent=2)}", "DEBUG", immediate=True)
        except Exception as e:
            log_process(f"Failed to parse response directly: {e}", "DEBUG", immediate=True)
            resume_data = clean_api_response(result)
            log_process(f"Cleaned resume data: {json.dumps(resume_data, indent=2)}", "DEBUG", immediate=True)
        if not isinstance(resume_data, dict):
            raise ResumeExtractionError("Invalid response format (not a dict).")
        required_keys = ["objective", "skills_list", "jobs_section", "education", "certifications"]
        for key in required_keys:
            if key not in resume_data:
                resume_data[key] = "" if key == "objective" else []
        resume_data.update({
            "rid": f"{datetime.datetime.now():%Y%m%d_%H%M%S}",
            "raw_text": raw_text,
            "source_file": file_name,
            "extraction_date": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "usage_count": 0
        })
        return ResumeData.from_dict(resume_data)
    except Exception as e:
        raise ResumeExtractionError(f"Failed to extract resume data: {e}")

def save_resume_data(resume_data: ResumeData) -> Path:
    """
    Saves extracted resume data as a JSON file.
    """
    output_dir = Path(CONFIG["EXTRACTED_DATA_DIR"]) / "resume_analysis"
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / f"RES-{resume_data.rid}.json"
    try:
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(vars(resume_data), f, indent=2)
        return output_path
    except Exception as e:
        raise ResumeExtractionError(f"Failed to save resume data: {e}")

def organize_resume_file(resume_file: Path, rid: str, processed_dir: Optional[Path] = None) -> None:
    """
    Moves processed resume file to the DONE directory with an RID prefix.
    """
    if processed_dir is None:
        processed_dir = Path(CONFIG["INPUT_RESUME"]) / "processed_resumes"
    done_dir = processed_dir / "DONE"
    try:
        done_dir.mkdir(parents=True, exist_ok=True)
        new_filename = f"RES-{rid}_{resume_file.name}"
        moved_path = done_dir / new_filename
        shutil.move(str(resume_file), str(moved_path))
    except Exception as e:
        raise ResumeExtractionError(f"Failed to organize resume file: {e}")

def process_resume_file(resume_file: FilePath) -> Optional[ResumeData]:
    """
    Processes a single resume file end-to-end.
    Skips processing if the resume has already been processed.
    """
    resume_file = Path(resume_file)
    if is_resume_file_processed(resume_file):
        log_process(f"Skipping already processed resume file: {resume_file.name}", "INFO")
        return None
    log_process(f"Processing resume file: {resume_file.name}", "INFO")
    try:
        raw_text = extract_text_from_file(resume_file)
        if not raw_text:
            raise ResumeExtractionError("Text extraction failed - empty result.")
        resume_data = extract_resume_data(raw_text, resume_file.name)
        json_path = save_resume_data(resume_data)
        organize_resume_file(resume_file, resume_data.rid)
        log_application({
            "date_time": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "JID": "",  # No job ID for resume extraction
            "RID": resume_data.rid,
            "short_company_name": "",
            "short_title": "",
            "job_folder": str(json_path.parent),
            "status": "ResumeExtracted",
            "match_score": 0,
            "final_resume_path": ""
        })
        return resume_data
    except Exception as e:
        log_process(f"Failed to process {resume_file.name}: {e}", "ERROR")
        return None

def process_resume_files(resume_file: FilePath) -> List[ResumeData]:
    """
    Process a single resume file and return the extracted data.
    
    Args:
        resume_file: Path to the resume file to process
        
    Returns:
        List containing the ResumeData if successful, empty list otherwise
    """
    if not isinstance(resume_file, Path):
        resume_file = Path(resume_file)
        
    if not resume_file.exists():
        log_process(f"Resume file not found: {resume_file}", "ERROR")
        return []
        
    if resume_file.suffix.lower() not in [".txt", ".pdf", ".docx"]:
        log_process(f"Unsupported file type: {resume_file.suffix}", "ERROR")
        return []
        
    try:
        result = process_resume_file(resume_file)
        return [result] if result else []
    except Exception as e:
        log_process(f"Failed to process {resume_file}: {e}", "ERROR")
        return []
