"""
resume_builder.py
2/19/2025
v8.0.4

This module assembles the final resume by injecting optimized content into a DOCX template.
Key features:

1. Template Processing:
   - Loads DOCX template with placeholders
   - Validates template structure
   - Handles placeholder replacement

2. Content Injection:
   - Overview/objective statement
   - Skills (exactly 10 skills)
   - Job bullet points with formatting
   - Maintains document styling

3. Output Generation:
   - Creates organized output directories
   - Validates final document
   - Logs detailed resume summaries
"""

import os
import json
import datetime
from pathlib import Path
from typing import Dict, List, Optional, Any, Union
from dataclasses import dataclass
from openai import OpenAI  


from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from logging_manager import log_process, log_resume_summary
from placeholder_matcher import PlaceholderMatcher
from config_manager import CONFIG, get_openai_api_key
from helpers import validate_file_path, safe_file_write

api_key = get_openai_api_key()  # Retrieve OpenAI API key
client = OpenAI(api_key=api_key)  # Initialize OpenAI client




# Type aliases
JSON = Dict[str, Any]
FilePath = Union[str, Path]

@dataclass
class OptimizedContent:
    """Container for optimized resume content"""
    objective: str
    skills: List[str]
    bullets: List[Dict[str, str]]
    match_rating: float
    job_data: Dict[str, str]
    
    @classmethod
    def from_dict(cls, data: Dict[str, Any]) -> 'OptimizedContent':
        """Creates an OptimizedContent instance from optimization results"""
        return cls(
            objective=data.get('new_objective', ''),
            skills=data.get('optimized_skills', '').split(', '),
            bullets=data.get('optimized_bullets', {}).get('bullets', []),
            match_rating=data.get('job_match_evaluation', {}).get('match_rating', 0),
            job_data=data.get('job_json', {})
        )
        
    def validate(self) -> None:
        """
        Validates the content structure.
        
        Raises:
            ValueError: If content is invalid
        """
        if not self.objective:
            raise ValueError("Missing objective statement")
            
        if len(self.skills) != 10:
            raise ValueError(f"Expected 10 skills, got {len(self.skills)}")
            
        required_job_fields = ['jid', 'Company Name', 'Title']
        missing = [f for f in required_job_fields if not self.job_data.get(f)]
        if missing:
            raise ValueError(f"Missing required job fields: {', '.join(missing)}")

class ResumeBuilderError(Exception):
    """Custom exception for resume building errors"""
    pass

def get_template_path() -> Path:
    """
    Gets the path to the resume template.
    
    Returns:
        Path: Template file path
        
    Raises:
        ResumeBuilderError: If template not found
    """
    template_path = Path(CONFIG["STATIC_DATA_DIR"]) / "resume_template" / "template-resume.docx"
    if not template_path.exists():
        raise ResumeBuilderError(f"Template not found at {template_path}")
    return template_path

def create_output_directory(
    job_data: Dict[str, str],
    base_dir: Optional[str] = None
) -> Path:
    """
    Creates and returns the output directory path.
    
    Args:
        job_data: Job metadata
        base_dir: Optional base directory override
        
    Returns:
        Path: Created directory path
        
    Raises:
        ResumeBuilderError: If directory creation fails
    """
    if base_dir is None:
        base_dir = CONFIG["FINISHED_JOB_RESUME_DIR"]
        
    try:
        today = datetime.datetime.now().strftime("%Y%m%d")
        company = job_data.get("Company Name", "Unknown").replace(" ", "")
        title = job_data.get("Title", "Unknown").replace(" ", "")
        jid = job_data.get("jid", "Unknown")
        
        output_dir = Path(base_dir) / f"{today}_{company}_{title}_{jid}"
        output_dir.mkdir(parents=True, exist_ok=True)
        return output_dir
        
    except Exception as e:
        raise ResumeBuilderError(f"Failed to create output directory: {e}")

def replace_placeholder_text(
    paragraph: Any,
    placeholder: str,
    new_text: str,
    preserve_formatting: bool = True
) -> None:
    """
    Replaces placeholder text while preserving formatting.
    
    Args:
        paragraph: Document paragraph
        placeholder: Text to replace
        new_text: Replacement text
        preserve_formatting: Whether to preserve run formatting
    """
    if not preserve_formatting:
        paragraph.text = paragraph.text.replace(placeholder, new_text)
        return
        
    for run in paragraph.runs:
        if placeholder in run.text:
            run.text = run.text.replace(placeholder, new_text)

def inject_content(
    doc: Document,
    content: OptimizedContent,
    skill_placeholders: List[str]
) -> None:
    """
    Injects optimized content into the document.
    
    Args:
        doc: Document object
        content: Optimized content
        skill_placeholders: List of skill placeholders
        
    Raises:
        ResumeBuilderError: If content injection fails
    """
    try:
        # Replace overview
        for para in doc.paragraphs:
            if "<OverView>" in para.text:
                replace_placeholder_text(para, "<OverView>", content.objective)
                
        # Replace skills
        for i, skill in enumerate(content.skills, 1):
            placeholder = f"<SKILL {i}>"
            for para in doc.paragraphs:
                if placeholder in para.text:
                    replace_placeholder_text(para, placeholder, skill)
                    
        # Replace bullet points
        for i, bullet in enumerate(content.bullets, 1):
            overview_placeholder = f"<Experience-Bullet{i}-BoldedOverview-J1>"
            detail_placeholder = f"<Experience-Bullet{i}-J1>"
            
            for para in doc.paragraphs:
                if overview_placeholder in para.text:
                    replace_placeholder_text(
                        para,
                        overview_placeholder,
                        bullet.get("bolded_overview", ""),
                        preserve_formatting=True
                    )
                if detail_placeholder in para.text:
                    replace_placeholder_text(
                        para,
                        detail_placeholder,
                        bullet.get("description", ""),
                        preserve_formatting=True
                    )
                    
    except Exception as e:
        raise ResumeBuilderError(f"Failed to inject content: {e}")

def log_resume_details(
    content: OptimizedContent,
    output_path: Path
) -> None:
    """
    Logs resume summary information.
    
    Args:
        content: Optimized content
        output_path: Path to generated resume
    """
    job_data = content.job_data
    field_full = job_data.get("field", "Unknown")
    
    log_resume_summary({
        "job_description_jid": job_data.get("jid", "Unknown"),
        "posting_date": job_data.get("posting_date", "Unknown"),
        "title": job_data.get("Title", "Unknown"),
        "location": job_data.get("Location", "Unknown"),
        "salary": job_data.get("Salary", "Unknown"),
        "field_full": field_full,
        "field_abbr": field_full[:2].upper() if field_full != "Unknown" else "UN",
        "final_resume_date": datetime.datetime.now().strftime("%Y-%m-%d"),
        "final_resume_match_percentage": content.match_rating,
        "final_resume_filename": output_path.name,
        "final_resume_link": str(output_path.resolve()),
        "submitted": "",
        "status": "Generated"
    })

def build_final_resume(
    optimized_data: Dict[str, Any],
    output_dir: Optional[str] = None
) -> Optional[Path]:
    """
    Builds the final resume from optimized content.
    
    Args:
        optimized_data: Optimization results
        output_dir: Optional output directory override
        
    Returns:
        Optional[Path]: Path to generated resume if successful
        
    Raises:
        ResumeBuilderError: If resume building fails
    """
    try:
        # Validate and structure input
        content = OptimizedContent.from_dict(optimized_data)
        content.validate()
        
        # Get template and create output directory
        template_path = get_template_path()
        output_dir = create_output_directory(content.job_data, output_dir)
        output_path = output_dir / "final_resume.docx"
        
        # Load and process template
        doc = Document(template_path)
        skill_placeholders = PlaceholderMatcher.extract_skill_placeholders(
            "\n".join(p.text for p in doc.paragraphs)
        )
        
        # Inject content and save
        inject_content(doc, content, skill_placeholders)
        doc.save(output_path)
        
        # Log details
        log_process(f"Final resume saved at {output_path}", "INFO")
        log_resume_details(content, output_path)
        
        return output_path
        
    except Exception as e:
        error_msg = f"Failed to build resume: {str(e)}"
        log_process(error_msg, "ERROR")
        raise ResumeBuilderError(error_msg)
